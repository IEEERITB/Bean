from docx import Document
import os

def create_template():
    # Ensure templates directory exists
    if not os.path.exists("templates"):
        os.makedirs("templates")
        
    doc = Document()
    
    doc.add_heading('IEEE Event Report', 0)
    
    doc.add_heading('Event Title:', level=1)
    doc.add_paragraph('{{ event_title }}')
    
    doc.add_heading('Date:', level=1)
    doc.add_paragraph('{{ date }}')
    
    doc.add_heading('Speaker:', level=1)
    doc.add_paragraph('{{ speaker_name }}')
    
    doc.add_heading('Attendance:', level=1)
    doc.add_paragraph('{{ attendance_count }}')
    
    doc.add_heading('Duration:', level=1)
    doc.add_paragraph('{{ duration_hours }} hours')
    
    doc.add_heading('Executive Summary:', level=1)
    doc.add_paragraph('{{ executive_summary }}')
    
    doc.add_heading('Key Takeaways:', level=1)
    # Jinja2 loop for list
    # syntax: {% for item in key_takeaways %} ... {% endfor %}
    # docxtpl handles this within the text
    doc.add_paragraph('{% for item in key_takeaways %}')
    doc.add_paragraph('- {{ item }}', style='List Bullet')
    doc.add_paragraph('{% endfor %}')
    
    save_path = "templates/ieee_report_template.docx"
    doc.save(save_path)
    print(f"Template saved to {save_path}")

if __name__ == "__main__":
    create_template()
